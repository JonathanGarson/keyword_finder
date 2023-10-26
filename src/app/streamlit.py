import streamlit as st
import pandas as pd
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR
import glob
import tempfile
import shutil

# The rest of your code remains the same...


def process_docx_files(docx_files, list_kw):
    """
    Go over all the provided DOCX files, highlight the keywords, extract the paragraph, and return a merged DOCX document.

    Args:
        docx_files: List of uploaded DOCX files.
        list_kw: List of keywords.

    Returns:
        Document: Merged DOCX document with highlighted paragraphs.
    """
    # Set up regex
    patterns = [r'\b' + word + r'\b' for word in list_kw]
    re_highlight = re.compile('(' + '|'.join(p for p in patterns) + ')+', re.IGNORECASE)

    merged_doc = Document()

    for docx_file in docx_files:
        with open(docx_file, "rb") as f: # Open the file in read-only mode
            title = docx_file.split("/")[-1]  # Extract the file name without extension
            print(title)
            doc = Document(docx_file)
            merged_doc.add_heading(title, level=1)  # Add the title as Header 1 to the merged document

            for para in doc.paragraphs:
                text = para.text
                if len(re_highlight.findall(text)) > 0:
                    matches = re_highlight.finditer(text)
                    p3 = 0
                    highlighted_para = merged_doc.add_paragraph()  # Add a new paragraph to the merged document
                    for match in matches:
                        p1 = p3
                        p2, p3 = match.span()
                        highlighted_para.add_run(text[p1:p2])
                        run = highlighted_para.add_run(text[p2:p3])
                        run.font.highlight_color = WD_COLOR.YELLOW
                        highlighted_para.add_run(text[p3:])

    return merged_doc

# def merge_all_docx(docx_files, output_path):
#     merged_doc = Document()
#     for docx_file in docx_files:
#         doc = Document(docx_file)
#         for element in doc.element.body:
#             merged_doc.element.body.append(element)
#     merged_doc.save(output_path)

def main():
    st.title("Keyword Finder App")
    st.write("Upload Word documents and an Excel file containing keywords")
    
    docx_files = st.file_uploader("Upload Word documents", type=["docx"], accept_multiple_files=True)
    xlsx_file = st.file_uploader("Upload Excel file", type=["xlsx"])

    if docx_files and xlsx_file:
        df_kw = pd.read_excel(xlsx_file)
        list_kw = df_kw.iloc[:, 0].tolist()

        # Create a temporary directory to store the uploaded DOCX files
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_file_list = []
            for docx_file in docx_files:
                # Save the uploaded DOCX files to the temporary directory
                temp_path = os.path.join(temp_dir, docx_file.name)
                with open(temp_path, 'wb') as f:
                    f.write(docx_file.read())
                docx_file_list.append(temp_path)
            print(docx_file_list)

            # Process the uploaded Word documents and get the merged DOCX document
            merged_doc = process_docx_files(docx_file_list, list_kw)
            with open(os.path.join(temp_dir, "final_merged.docx"), "wb") as f:
                merged_doc.save(f)

            # merged_filename = st.text_input("Merged DOCX Filename", "final_merged")
                
            # if st.button("Download Merged DOCX"):
            #     st.markdown("### Download the merged DOCX file")
            #     with st.spinner("Merging and saving..."):
            #         # Merge the docx files and save the result
            #         merge_all_docx(docx_file_list, os.path.join(temp_dir, merged_filename + ".docx"))

            #         # Move the merged file to the user-specified location
            output_path = os.path.join(temp_dir, "final_merged.docx")
            #         shutil.move(output_path, ".")

            #     st.success("Merged DOCX file has been saved!")


            with open(output_path, "rb") as file:
                btn = st.download_button(
                        label="Download docx",
                        data=file,
                        file_name="final_merged.docx",
                        mime="application/octet-stream"
                    )

            st.success("Merged DOCX file has been saved!")

if __name__ == "__main__":
    main()
