import pandas as pd
import glob
import re
import os
from docx import Document
from docx.enum.text import WD_COLOR

# Define a function to extract the file name without extension
def extract_file_name(file: str):
    """
    Extract the file name without extension

    Args:
        file: The file path

    Returns:
        The file name (str) without extension
    """
    file_name = os.path.basename(file)
    file_name_without_extension = os.path.splitext(file_name)[0]
    return file_name_without_extension

# Define a function to process all *.docx files in a folder
def process_docx_files(input_folder_path:str, output_folder_path:str):
    """
    Go over all the files in a directory, highlight the keywords, extract the paragraph and then save the modified paragraph to a new document

    Args:
        input_folder_path: The path to the folder containing the *.docx files
        output_folder_path: The path to the folder to store the modified *.docx files
    """
    # Set up regex
    patterns = [r'\b' + word + r'\b' for word in list_kw]
    re_highlight = re.compile('(' + '|'.join(p for p in patterns) + ')+', re.IGNORECASE)

    # Iterate over all the *.docx files in the folder
    docx_files = glob.glob(input_folder_path + r'/*.docx')
    try:
        for file in docx_files:
            doc = Document(file)
            title = extract_file_name(file)  # Extract the file name without extension
            modified_doc = Document()  # Create a new document to store modified paragraphs
            modified_doc.add_heading(title, level=1)  # Add the title as Header 1 to the modified document
            for para in doc.paragraphs:
                text = para.text
                if len(re_highlight.findall(text)) > 0:
                    matches = re_highlight.finditer(text)
                    p3 = 0
                    highlighted_para = modified_doc.add_paragraph()  # Add a new paragraph to the modified document
                    for match in matches:
                        p1 = p3
                        p2, p3 = match.span()
                        highlighted_para.add_run(text[p1:p2])
                        run = highlighted_para.add_run(text[p2:p3])
                        run.font.highlight_color = WD_COLOR.YELLOW
                        highlighted_para.add_run(text[p3:])
            if modified_doc.paragraphs:  # Only save the modified document if it contains highlighted paragraphs
                modified_doc.save(os.path.join(output_folder_path, rf"{title}.docx"))
    except Exception as e:
        print(f"Error processing {file}: {e}")

def merge_all_docx(input_folder_path:str, output_folder_path:str):
    """
    Merge all the docx files in a folder into a single docx file

    Args:
        input_folder_path: The path to the folder containing the *.docx files
        output_folder_path: The path to the folder to store the merged *.docx file
    """
    docx_files = glob.glob(input_folder_path + r'/*.docx')
    documents = [Document(docx) for docx in docx_files]
    merged_document = Document()
    for doc in documents:
        for element in doc.element.body:
            merged_document.element.body.append(element)
    merged_document.save(output_folder_path + r'/final_merged.docx')

if __name__ == '__main__':

    # Import keywords from the Excel file and turn them into a list
    df_kw = pd.read_excel(r'./data/raw/keywords.xlsx')
    df_kw.to_csv(r'./data/raw/keywords.csv', index=False)
    list_kw = df_kw.iloc[:, 0].tolist()

    # Extract and highlight keywords in the docx files
    input_folder_path = r'./data/text/raw_docx'
    output_folder_path = r'./data/text/treated_docx'
    process_docx_files(input_folder_path=input_folder_path, output_folder_path= output_folder_path)

    # Merged the docx files into a single docx file
    input_folder_merged = r'./data/text/treated_docx'
    output_folder_merged = r'./data/text'
    merge_all_docx(input_folder_path=input_folder_merged, output_folder_path=output_folder_merged)